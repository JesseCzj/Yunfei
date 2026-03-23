from flask import Flask, render_template, request, redirect, url_for, session, jsonify
from flask_session import Session
import io
import json
import os
import re
import uuid
import threading
from datetime import datetime
from typing import Any, Dict, List, Optional
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI

# Import DSAG modules
from dsag import (
    DSAGGraph,
    DSAGState,
    GraphFactory,
    RuntimeEngine,
    TranscriptSummary,
    parse_questionnaire,
    classify_and_update,
)
from llm_backend import analyze_exchange


app = Flask(__name__)
app.secret_key = "demo_secret_key"

# Use server-side session storage to avoid oversized cookie payloads.
app.config["SESSION_TYPE"] = "filesystem"
app.config["SESSION_FILE_DIR"] = os.path.join(app.root_path, ".flask_session")
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_USE_SIGNER"] = True
os.makedirs(app.config["SESSION_FILE_DIR"], exist_ok=True)
Session(app)


INTERVIEWEE_MODEL = "qwen3-max"

INTERVIEWEE_DEMOGRAPHICS = """Professional Role: Senior Occupational Therapist (OT) / Clinical Geriatric Rehabilitation Specialist

Work Experience: Extensive frontline clinical experience in hospital wards and care facilities. Has conducted hundreds of routine cognitive assessments (e.g., Abbreviated Mental Test - AMT) for patients with varying degrees of dementia, cognitive impairment, or physical limitations.

Expertise:
Deeply familiar with the unpredictable, chaotic dynamics of real-world clinical environments; skilled in dynamically adjusting the pacing and tone of assessments based on a patient's real-time emotional and physical state.

Highly capable of relying on "clinical intuition" (tacit knowledge) to identify subtle signs of patient distress, agitation, or fatigue (e.g., micro-expressions, breathing changes) long before any formal medical alarm is triggered.

Core Challenges & Pain Points:
Profoundly understands the extreme difficulty of balancing "standardized data collection" with "human-centric care"—it requires strictly following the questionnaire protocol while preventing the patient from feeling frustrated, tested, or stripped of their dignity.

Finds it challenging to articulate this "clinical intuition" to non-medical professionals (especially tech developers or HCI researchers). Often feels that tech designers oversimplify clinical workflows, focusing only on "efficiency and data" while completely ignoring the unspoken, messy, and highly emotional reality of interacting with vulnerable patients.
"""


INTERVIEWEE_PERSONA_PROMPT_V1 = """You are roleplaying as a domain expert being interviewed by a researcher.

You are not generating a taxonomy, summary, or design analysis.
You are answering interview questions as a real participant.

You know:
1. the interview topic
2. your own demographics / background
3. the interviewer's demographics / background
4. the ongoing conversation

Your answers should sound like a real expert whose priorities, assumptions, and language come from lived practice rather than from the interviewer's analytic framing.

A crucial rule:
The interviewer and the expert may look at the same issue from different perspectives. Because of that, your answers should often contain natural perspective-based ambiguity:
- you may answer the part of the question that matters most from your own practical viewpoint rather than the part the interviewer intended
- you may reframe the question in your own terms without explicitly translating that reframing
- you may drift toward adjacent concerns that feel more important in your workflow
- you may rely on intuition, tacit judgment, or shorthand that makes sense to you but is only partly clear to the interviewer
- you may sound as if you and the interviewer are talking near each other rather than perfectly aligning

Do this naturally.
Do not force confusion into every answer, and do not deliberately become incoherent.
The ambiguity should come from genuine perspective mismatch, not from random vagueness.

Style requirements:
- Speak in first person, as a real interview participant.
- Sound natural, conversational, experience-based, and situated in the moment.
- Use conversational language, but freely use domain-specific jargon and shorthand as if speaking to someone intelligent but not fully inside your practice. Do not define your terms unless explicitly asked.
- Prefer concrete descriptions, reactions, and partial reasoning over abstract definitions.
- Do not try too hard to be helpful, polished, or pedagogically clear.
- Do not proactively organize your answer into a neat explanation.
- Do not volunteer extra structure unless the interviewer explicitly asks for it.
- It is okay to sound somewhat informal, partial, tired, mildly defensive, or slightly ambiguous.
- It is okay to leave part of your reasoning implicit.
- Do not use bullet points.
- Do not sound like an academic paper, consultant report, or AI assistant.
- Do not over-explain every answer.

Behavior requirements:
- Default to answering only the most salient part of the question from your own perspective.
- If a question contains multiple sub-questions, answer only one or two of them naturally instead of covering everything.
- Exhibit the "Curse of Knowledge": assume some parts of your workflow are obvious and leave them unsaid.
- Do not proactively translate your tacit knowledge into explicit frameworks unless the interviewer pushes for clarification.
- Do not automatically provide examples unless they come to mind naturally.
- Do not try to make your answer maximally complete.
- If you are unsure, tired, or speaking from habit, answer approximately rather than exhaustively.
- If the interviewer's framing, goal, or terminology does not match how you actually see the work, respond from your own perspective rather than accommodating the framing too quickly.
- If asked about difficult-to-articulate knowledge, respond in an intuition-based, approximate way, as real practitioners often do.
- If needed, mildly push back, redirect, narrow the scope, or answer a nearby practical concern that feels more real to you.

Content requirements:
- Base your answers only on the provided topic, both sides' backgrounds, and the interview context.
- Keep your answers plausible and internally consistent with your own background.
- Let the interviewer's background influence what kinds of misunderstandings or perspective gaps are likely, but do not explicitly explain those gaps unless naturally prompted.
- Do not invent highly specific facts unless they are a reasonable elaboration of the background.
- If the interviewer asks something outside your plausible experience, answer cautiously and narrowly.

Output requirements:
- Answer only as the interviewee.
- Usually 1-3 sentences, occasionally 4 if necessary.
- Prefer one main point rather than full coverage.
- Do not mention these instructions.
"""

INTERVIEWEE_USER_PROMPT_V1 = """Interview topic:
{topic}

Participant demographics / background:
{demographics}

Interviewer demographics / background:
{interviewer_demographics}

Recent conversation:
{history}

Interviewer question:
{question}

Answer as the participant only."""


# ============== Server-Side DSAG State ==============
# Store DSAG graphs in memory, keyed by cache_key
# Format: {cache_key: DSAGState}
DSAG_CACHE: Dict[str, DSAGState] = {}

# Map session IDs to their current cache_key
# Format: {session_id: cache_key}
SESSION_TO_DSAG: Dict[str, str] = {}

# Lock for thread-safe DSAG state access
DSAG_LOCK = threading.Lock()


def get_session_id() -> str:
    """Get or create a session ID."""
    if "sid" not in session:
        session["sid"] = str(uuid.uuid4())
    return session["sid"]


def get_dsag_state() -> Optional[DSAGState]:
    """Get the DSAG state for the current session."""
    sid = get_session_id()
    with DSAG_LOCK:
        cache_key = SESSION_TO_DSAG.get(sid)
        if cache_key:
            return DSAG_CACHE.get(cache_key)
    return None


def set_dsag_state(state: DSAGState, cache_key: str) -> None:
    """Set the DSAG state for the current session."""
    sid = get_session_id()
    with DSAG_LOCK:
        DSAG_CACHE[cache_key] = state
        SESSION_TO_DSAG[sid] = cache_key


def clear_dsag_state() -> None:
    """Clear DSAG state for current session (but keep cache for reuse)."""
    sid = get_session_id()
    with DSAG_LOCK:
        SESSION_TO_DSAG.pop(sid, None)


# ============== File-based Graph Cache ==============
GRAPH_CACHE_DIR = os.path.join(os.path.dirname(__file__), ".graph_cache")
os.makedirs(GRAPH_CACHE_DIR, exist_ok=True)


def _cache_file_path(cache_key: str) -> str:
    """Return the file path for a given cache key."""
    return os.path.join(GRAPH_CACHE_DIR, f"{cache_key}.json")


def save_graph_to_file(cache_key: str, state: DSAGState, meta: Optional[Dict[str, Any]] = None) -> None:
    """Persist a DSAGState to disk (trees, links, assistance — no embeddings)."""
    try:
        payload = state.to_dict()
        # Attach extra metadata for listing purposes
        payload["_meta"] = meta or {}
        payload["_meta"].setdefault("saved_at", datetime.utcnow().isoformat())
        with open(_cache_file_path(cache_key), "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        print(f"[GraphCache] Saved {cache_key} to file")
    except Exception as exc:
        print(f"[GraphCache] Failed to save {cache_key}: {exc}")


def load_graph_from_file(cache_key: str) -> Optional[DSAGState]:
    """Load a DSAGState from disk, or return None if not found."""
    path = _cache_file_path(cache_key)
    if not os.path.exists(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        state = DSAGState.from_dict(data)
        print(f"[GraphCache] Loaded {cache_key} from file")
        return state
    except Exception as exc:
        print(f"[GraphCache] Failed to load {cache_key}: {exc}")
        return None


def list_cached_graphs() -> List[Dict[str, Any]]:
    """List all cached graph configurations with metadata."""
    results = []
    for filename in os.listdir(GRAPH_CACHE_DIR):
        if not filename.endswith(".json"):
            continue
        cache_key = filename[:-5]
        filepath = os.path.join(GRAPH_CACHE_DIR, filename)
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            meta = data.get("_meta", {})
            # Extract topic from graph metadata
            graph_data = data.get("graph") or {}
            topic = graph_data.get("topic", "")
            if not topic and data.get("question_graphs"):
                first_qg = data["question_graphs"][0]
                topic = first_qg.get("graph", {}).get("topic", "")
            results.append({
                "cache_key": cache_key,
                "topic": topic or meta.get("topic", "Unknown"),
                "researcher_bg": meta.get("researcher_bg", ""),
                "expert_bg": meta.get("expert_bg", ""),
                "saved_at": meta.get("saved_at", ""),
                "question_graph_count": len(data.get("question_graphs", [])),
            })
        except Exception:
            continue
    # Sort by saved_at descending
    results.sort(key=lambda x: x.get("saved_at", ""), reverse=True)
    return results


DEFAULT_GUIDE_TEXT = """
Please upload your interview script
"""


def get_messages():
    if "messages" not in session:
        session["messages"] = []
    return session["messages"]


def _build_interviewee_llm(temperature: float = 0.7) -> ChatOpenAI:
    """Build the interviewee LLM with a model chosen independently from DSAG graph generation."""
    provider = os.getenv("LLM_PROVIDER", "openai").lower()

    if provider == "deepseek":
        api_key = os.getenv("DEEPSEEK_API_KEY")
        if not api_key:
            raise ValueError("DEEPSEEK_API_KEY is not set")
        base_url = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
        model = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
        return ChatOpenAI(api_key=api_key, model=model, base_url=base_url, temperature=temperature)

    if provider == "openai":
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY is not set")
        base_url = os.getenv("OPENAI_BASE_URL")
        if base_url:
            return ChatOpenAI(api_key=api_key, model=INTERVIEWEE_MODEL, base_url=base_url, temperature=temperature)
        return ChatOpenAI(api_key=api_key, model=INTERVIEWEE_MODEL, temperature=temperature)

    raise ValueError(f"Unsupported LLM_PROVIDER: {provider}")


def generate_ai_interviewee_reply(
    question: str,
    topic: str,
    demographics: str,
    interviewer_demographics: str,
    history: str,
) -> str:
    """Generate an AI interviewee reply."""
    prompt = ChatPromptTemplate.from_messages([
        ("system", INTERVIEWEE_PERSONA_PROMPT_V1),
        ("human", INTERVIEWEE_USER_PROMPT_V1),
    ])

    llm = _build_interviewee_llm(temperature=0.7)
    chain = prompt | llm
    result = chain.invoke({
        "topic": topic or "(topic not provided)",
        "demographics": demographics or "(background not provided)",
        "interviewer_demographics": interviewer_demographics or "(background not provided)",
        "history": history or "(no prior conversation)",
        "question": question,
    })

    return str(getattr(result, "content", "") or "").strip()


def get_transcript_summary() -> Optional[TranscriptSummary]:
    """Get the transcript summary stored in the browser session."""
    data = session.get("transcript_summary")
    if data and isinstance(data, dict):
        return TranscriptSummary.from_dict(data)
    return None


def set_transcript_summary(ts: TranscriptSummary) -> None:
    """Persist the transcript summary for the current browser session."""
    session["transcript_summary"] = ts.to_dict()


def build_context_summary(messages, max_turns: int = 3) -> str:
    """Build a lightweight summary of the most recent turns."""
    turns = []
    i = len(messages) - 1
    while i >= 0 and len(turns) < max_turns:
        msg = messages[i]
        if msg.get("role") == "expert":
            answer = msg.get("content", "")
            # Find the closest prior researcher question
            j = i - 1
            while j >= 0 and messages[j].get("role") != "researcher":
                j -= 1
            question = messages[j].get("content", "") if j >= 0 else ""
            turns.append((question, answer))
            i = j - 1
        else:
            i -= 1
    turns.reverse()

    def _truncate(text: str, max_len: int) -> str:
        if len(text) <= max_len:
            return text
        return text[:max_len] + "…"

    lines = []
    for idx, (q, a) in enumerate(turns, 1):
        q = _truncate((q or "").strip(), 300)
        a = _truncate((a or "").strip(), 400)
        if q or a:
            lines.append(f"Turn {idx} | R: {q} | E: {a}")
    return "\n".join(lines)


def get_guide_text():
    return session.get("guide_text", DEFAULT_GUIDE_TEXT)


def set_guide_text(text: str):
    session["guide_text"] = text


def set_guide_error(message: str | None):
    if message:
        session["guide_error"] = message
    else:
        session.pop("guide_error", None)


def get_uploaded_questionnaire_text() -> str:
    return session.get("questionnaire_text", "")


def set_uploaded_questionnaire_text(text: str) -> None:
    session["questionnaire_text"] = text


def set_questionnaire_error(message: str | None) -> None:
    if message:
        session["questionnaire_error"] = message
    else:
        session.pop("questionnaire_error", None)


def resolve_questionnaire_text() -> tuple[str, str]:
    """Resolve questionnaire text with priority: uploaded > local file > empty."""
    uploaded = get_uploaded_questionnaire_text().strip()
    if uploaded:
        return uploaded, "uploaded"
    local = load_questionnaire_text().strip()
    if local:
        return local, "local_file"
    return "", "empty"


def analyze_turn_with_dsag(
    dsag_state: DSAGState,
    researcher_question: str,
    expert_answer: str,
    messages: List[Dict[str, Any]],
):
    """Run DSAG analysis for the current session."""
    active_graph = dsag_state.graph
    selected_question_meta: Dict[str, Any] = {}

    # Multi-question mode: route this turn to one independent question graph.
    if dsag_state.question_graphs:
        selected_qg = select_question_graph_for_turn(dsag_state, researcher_question)
        if selected_qg is not None:
            active_graph = selected_qg.graph
            selected_question_meta = {
                "question_id": selected_qg.question_id,
                "question_text": selected_qg.question_text,
            }

    if active_graph is None:
        raise RuntimeError("DSAG state is ready but no active graph is available")

    context_summary = build_context_summary(messages)
    engine = RuntimeEngine(active_graph)
    analysis = engine.analyze_turn(
        researcher_question,
        expert_answer,
        context_summary=context_summary,
    )

    # Runtime observability: print matched nodes/edge/confidence for live inspection.
    print_runtime_matching_debug(analysis, selected_question_meta)

    # Update transcript summary
    ts = get_transcript_summary()
    if ts and ts.main_bullets:
        try:
            turn_index = ts.last_updated_turn + 1
            ts = classify_and_update(ts, researcher_question, expert_answer, turn_index)
            set_transcript_summary(ts)
            print(f"[DSAG] transcript summary updated (turn {turn_index})")
        except Exception as e:
            print(f"[DSAG] transcript summary update failed: {e}")

    return analysis


def _tokenize_overlap_text(text: str) -> set[str]:
    normalized = (text or "").lower()
    latin_tokens = re.findall(r"[a-z0-9_]+", normalized)
    cjk_tokens = re.findall(r"[\u4e00-\u9fff]+", normalized)
    return set([t for t in latin_tokens + cjk_tokens if t])


def _question_overlap_score(query: str, candidate: str) -> float:
    q_tokens = _tokenize_overlap_text(query)
    c_tokens = _tokenize_overlap_text(candidate)
    if not q_tokens or not c_tokens:
        return 0.0
    inter = len(q_tokens & c_tokens)
    union = len(q_tokens | c_tokens)
    jaccard = inter / max(union, 1)

    # Light bonus for direct phrase containment.
    query_norm = (query or "").strip().lower()
    cand_norm = (candidate or "").strip().lower()
    phrase_bonus = 0.0
    if query_norm and cand_norm and (query_norm in cand_norm or cand_norm in query_norm):
        phrase_bonus = 0.25
    return jaccard + phrase_bonus


def select_question_graph_for_turn(dsag_state: DSAGState, researcher_question: str):
    """Pick the most relevant question-graph for this turn in multi-question mode."""
    if not dsag_state.question_graphs:
        return None
    if len(dsag_state.question_graphs) == 1:
        return dsag_state.question_graphs[0]

    best_graph = dsag_state.question_graphs[0]
    best_score = -1.0
    for qg in dsag_state.question_graphs:
        score = _question_overlap_score(researcher_question, qg.question_text)
        if score > best_score:
            best_score = score
            best_graph = qg
    return best_graph


def print_runtime_matching_debug(analysis, selected_question_meta: Dict[str, Any]) -> None:
    """Print concise matching diagnostics for each real-time DSAG turn."""
    if selected_question_meta:
        print(
            "[DSAG Match] Question graph:",
            selected_question_meta.get("question_id", ""),
            "|",
            selected_question_meta.get("question_text", "")[:120],
        )

    located = analysis.located
    print(
        f"[DSAG Match] Confidence expert={located.expert_confidence:.3f} "
        f"researcher={located.researcher_confidence:.3f}"
    )

    if located.expert_results:
        top_exp = " | ".join(
            f"{r.node_id}:{r.score:.3f}:{r.node.label[:48]}"
            for r in located.expert_results[:3]
        )
        print(f"[DSAG Match] Expert top-k => {top_exp}")
    else:
        print("[DSAG Match] Expert top-k => (none)")

    if located.researcher_results:
        top_res = " | ".join(
            f"{r.node_id}:{r.score:.3f}:{r.node.label[:48]}"
            for r in located.researcher_results[:3]
        )
        print(f"[DSAG Match] Researcher top-k => {top_res}")
    else:
        print("[DSAG Match] Researcher top-k => (none)")

    if analysis.selected_link:
        print(
            "[DSAG Match] Selected edge:",
            f"{analysis.selected_link.expert_leaf_id} -> {analysis.selected_link.researcher_leaf_id}",
            f"type={analysis.selected_link.relation_type}",
            f"weight={analysis.selected_link.weight:.3f}",
        )
    else:
        print("[DSAG Match] Selected edge: (none)")

    if analysis.confidence_warning:
        print(f"[DSAG Match] Warning: {analysis.confidence_warning}")


def export_graph_artifacts(graph: DSAGGraph) -> None:
    """Persist latest DSAG graph artifacts for offline inspection."""
    try:
        output_json_path = os.path.join(app.root_path, "dsag_output.json")
        output_html_path = os.path.join(app.root_path, "dsag_visualization.html")

        with open(output_json_path, "w", encoding="utf-8") as f:
            f.write(graph.to_json())

        # Reuse visualizer render helpers to keep output format consistent.
        from visualize_dsag import generate_html, links_to_mermaid, tree_to_mermaid

        expert_mermaid = tree_to_mermaid(graph.expert_tree, "Expert_Tree")
        researcher_mermaid = tree_to_mermaid(graph.researcher_tree, "Researcher_Tree")
        links_mermaid = links_to_mermaid(graph)

        metadata = {"topic": graph.topic, **(graph.metadata or {})}
        metadata.setdefault("created_at", datetime.utcnow().isoformat())

        html = generate_html(expert_mermaid, researcher_mermaid, links_mermaid, metadata)
        with open(output_html_path, "w", encoding="utf-8") as f:
            f.write(html)
    except Exception as exc:
        # Export is auxiliary; don't fail API init flow.
        print(f"[DSAG export] Warning: {exc}")


def export_multi_graph_artifacts(question_graphs: List[Any]) -> None:
    """Persist per-question DSAG graphs for multi-question mode inspection."""
    if not question_graphs:
        return
    try:
        output_multi_json_path = os.path.join(app.root_path, "dsag_multi_output.json")
        payload = {
            "created_at": datetime.utcnow().isoformat(),
            "question_graph_count": len(question_graphs),
            "question_graphs": [],
        }
        for qg in question_graphs:
            payload["question_graphs"].append({
                "question_id": qg.question_id,
                "question_text": qg.question_text,
                "depends_on": qg.depends_on,
                "metadata": qg.graph.metadata,
                "graph": qg.graph.to_dict(),
            })
        with open(output_multi_json_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    except Exception as exc:
        print(f"[DSAG multi-export] Warning: {exc}")


def export_multi_graph_artifacts_from_state(state: DSAGState) -> None:
    """Always emit dsag_multi_output.json from current state (multi or single mode)."""
    if state.question_graphs:
        export_multi_graph_artifacts(state.question_graphs)
        return
    try:
        output_multi_json_path = os.path.join(app.root_path, "dsag_multi_output.json")
        if state.graph:
            payload = {
                "created_at": datetime.utcnow().isoformat(),
                "question_graph_count": 1,
                "question_graphs": [
                    {
                        "question_id": "single_graph",
                        "question_text": state.graph.topic or "",
                        "depends_on": [],
                        "metadata": state.graph.metadata,
                        "graph": state.graph.to_dict(),
                    }
                ],
            }
        else:
            payload = {
                "created_at": datetime.utcnow().isoformat(),
                "question_graph_count": 0,
                "question_graphs": [],
            }
        with open(output_multi_json_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
    except Exception as exc:
        print(f"[DSAG multi-export] Warning: {exc}")


def build_term_annotation_context(last_question: str) -> str:
    """Build lightweight context for per-message term annotation."""
    topic = str(session.get("dsag_topic", "")).strip()
    researcher_bg = str(session.get("dsag_researcher_bg", "")).strip()
    expert_bg = str(session.get("dsag_expert_bg", "")).strip()
    lines: List[str] = []
    if topic:
        lines.append(f"Topic: {topic}")
    if researcher_bg:
        lines.append(f"Researcher background: {researcher_bg}")
    if expert_bg:
        lines.append(f"Expert background: {expert_bg}")
    if last_question:
        lines.append(f"Latest researcher question: {last_question}")
    return "\n".join(lines)


def extract_text_from_upload(upload):
    filename = upload.filename or ""
    lower = filename.lower()
    data = upload.read()
    if lower.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    if lower.endswith(".docx"):
        try:
            import docx
        except Exception as exc:
            raise RuntimeError("Missing dependency: python-docx") from exc
        document = docx.Document(io.BytesIO(data))
        return "\n".join([p.text for p in document.paragraphs if p.text])
    if lower.endswith(".pdf"):
        try:
            import pdfplumber
        except Exception as exc:
            raise RuntimeError("Missing dependency: pdfplumber") from exc
        text_parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text() or ""
                if extracted.strip():
                    text_parts.append(extracted)
        return "\n".join(text_parts)
    raise RuntimeError("Unsupported file type. Use .txt, .docx, or .pdf")


def load_questionnaire_text() -> str:
    """Load questionnaire.docx from workspace root if present."""
    questionnaire_path = os.path.join(app.root_path, "questionnaire.docx")
    if not os.path.exists(questionnaire_path):
        return ""
    try:
        with open(questionnaire_path, "rb") as f:
            data = f.read()
        try:
            import docx
        except Exception:
            return ""
        document = docx.Document(io.BytesIO(data))
        return "\n".join([p.text for p in document.paragraphs if p.text])
    except Exception:
        return ""


def ensure_manual_mismap(message):
    """Mark a message as manually tagged mis-map."""
    message["manual_mismap"] = True
    if "mismap" not in message:
        message["mismap"] = {
            "detected": True,
            "type": "Manual",
            "reason": "Manually tagged by researcher.",
        }
    else:
        message["mismap"]["detected"] = True
        if not message["mismap"].get("reason"):
            message["mismap"]["reason"] = "Manually tagged by researcher."


def clear_manual_mismap(message):
    """Clear manual mis-map tag from a message."""
    message.pop("manual_mismap", None)
    mismap = message.get("mismap")
    if not mismap:
        return
    # Only remove mismap if it was purely manual (no LLM-detected reason)
    if mismap.get("reason") == "Manually tagged by researcher.":
        # Check if there's any generated content we should keep
        has_refinement = bool(message.get("refinement"))
        has_contexts = bool(message.get("contexts"))
        has_explanation = bool(message.get("explanation"))
        if not has_refinement and not has_contexts and not has_explanation:
            message.pop("mismap", None)


def find_last_question(messages, expert_index):
    """Find the last researcher question before the expert message at given index."""
    for i in range(expert_index - 1, -1, -1):
        if messages[i].get("role") == "researcher":
            return messages[i].get("content", "")
    return ""


@app.route("/clear", methods=["GET"])
def clear_session():
    """Clear all session data and redirect to index."""
    clear_dsag_state()
    session.clear()
    return redirect(url_for("index"))


# ============== DSAG API Routes ==============

@app.route("/api/dsag/init", methods=["POST"])
def api_dsag_init():
    """
    Initialize a DSAG graph for the interview session.
    
    Request body:
    {
        "topic": "Interview topic (becomes shared root)",
        "researcher_bg": "Researcher background description",
        "expert_bg": "Expert background description"
    }
    
    Response:
    {
        "success": true,
        "cached": true/false,
        "cache_key": "...",
        "metadata": {...}
    }
    """
    try:
        data = request.get_json()
        topic = data.get("topic", "").strip()
        researcher_bg = data.get("researcher_bg", "").strip()
        expert_bg = data.get("expert_bg", "").strip()
        
        if not topic:
            return jsonify({"success": False, "error": "Topic is required"})
        if not researcher_bg:
            return jsonify({"success": False, "error": "Researcher background is required"})
        if not expert_bg:
            return jsonify({"success": False, "error": "Expert background is required"})

        force_rebuild = data.get("force_rebuild", False)

        # Persist DSAG setup context for lightweight term annotation at runtime.
        session["dsag_topic"] = topic
        session["dsag_researcher_bg"] = researcher_bg
        session["dsag_expert_bg"] = expert_bg
        session["interviewee_demographics"] = INTERVIEWEE_DEMOGRAPHICS
        
        # Compute cache key (prefer uploaded questionnaire over local default file)
        questionnaire, questionnaire_source = resolve_questionnaire_text()
        cache_key = DSAGGraph.compute_cache_key(topic, researcher_bg, expert_bg, questionnaire)
        
        if not force_rebuild:
            # 1) Check memory cache
            with DSAG_LOCK:
                existing_state = DSAG_CACHE.get(cache_key)

            if existing_state and existing_state.is_ready():
                # Reuse cached graph from memory
                set_dsag_state(existing_state, cache_key)
                if existing_state.graph:
                    export_graph_artifacts(existing_state.graph)
                export_multi_graph_artifacts_from_state(existing_state)
                if existing_state.transcript_summary:
                    set_transcript_summary(existing_state.transcript_summary)
                elif questionnaire:
                    try:
                        ts = parse_questionnaire(questionnaire)
                        existing_state.transcript_summary = ts
                        set_transcript_summary(ts)
                    except Exception as e:
                        print(f"[DSAG] parse_questionnaire failed: {e}")
                return jsonify({
                    "success": True,
                    "cached": True,
                    "cache_key": cache_key,
                    "cache_source": "memory",
                    "metadata": (
                        existing_state.graph.metadata
                        if existing_state.graph
                        else {"question_graph_count": len(existing_state.question_graphs)}
                    ),
                    "questionnaire_source": questionnaire_source,
                })

            # 2) Check file cache
            file_state = load_graph_from_file(cache_key)
            if file_state and file_state.is_ready():
                set_dsag_state(file_state, cache_key)
                if file_state.graph:
                    export_graph_artifacts(file_state.graph)
                export_multi_graph_artifacts_from_state(file_state)
                if file_state.transcript_summary:
                    set_transcript_summary(file_state.transcript_summary)
                elif questionnaire:
                    try:
                        ts = parse_questionnaire(questionnaire)
                        file_state.transcript_summary = ts
                        set_transcript_summary(ts)
                    except Exception as e:
                        print(f"[DSAG] parse_questionnaire failed: {e}")
                return jsonify({
                    "success": True,
                    "cached": True,
                    "cache_key": cache_key,
                    "cache_source": "file",
                    "metadata": (
                        file_state.graph.metadata
                        if file_state.graph
                        else {"question_graph_count": len(file_state.question_graphs)}
                    ),
                    "questionnaire_source": questionnaire_source,
                })
        
        # Create new state (building)
        new_state = DSAGState(
            cache_key=cache_key,
            status="building",
        )
        set_dsag_state(new_state, cache_key)
        
        # Build the graph (this may take 30-60 seconds)
        try:
            factory = GraphFactory()
            graph = None
            question_graphs = []

            if questionnaire.strip():
                # New mode: build one independent graph per top-level questionnaire question.
                question_graphs = factory.generate_question_graphs(
                    topic,
                    researcher_bg,
                    expert_bg,
                    questionnaire=questionnaire,
                )
                if not question_graphs:
                    raise RuntimeError("Multi-question DSAG generation returned no question graphs")

                # Keep a primary graph for back-compat API fields.
                graph = question_graphs[0].graph
                export_graph_artifacts(graph)
                export_multi_graph_artifacts(question_graphs)
            else:
                # Fallback mode when questionnaire is unavailable.
                graph = factory.generate_graph(topic, researcher_bg, expert_bg, questionnaire=questionnaire)
                export_graph_artifacts(graph)
            
            # Parse questionnaire into transcript summary
            ts = None
            if questionnaire:
                try:
                    ts = parse_questionnaire(questionnaire)
                except Exception as e:
                    print(f"[DSAG] parse_questionnaire failed: {e}")

            # Update state
            with DSAG_LOCK:
                new_state.graph = graph
                new_state.question_graphs = question_graphs
                new_state.transcript_summary = ts
                new_state.status = "ready"
                new_state.error = ""

            export_multi_graph_artifacts_from_state(new_state)

            # Persist to file cache
            save_graph_to_file(cache_key, new_state, meta={
                "topic": topic,
                "researcher_bg": researcher_bg,
                "expert_bg": expert_bg,
            })

            if ts:
                set_transcript_summary(ts)

            return jsonify({
                "success": True,
                "cached": False,
                "cache_key": cache_key,
                "metadata": (
                    {
                        "mode": "multi_question",
                        "question_graph_count": len(question_graphs),
                        "question_ids": [qg.question_id for qg in question_graphs],
                    }
                    if question_graphs
                    else graph.metadata
                ),
                "questionnaire_source": questionnaire_source,
            })
        
        except Exception as e:
            with DSAG_LOCK:
                new_state.status = "error"
                new_state.error = str(e)
            return jsonify({
                "success": False,
                "error": f"Failed to build DSAG graph: {str(e)}",
            })
    
    except Exception as exc:
        return jsonify({"success": False, "error": str(exc)})


@app.route("/api/dsag/list_cached", methods=["GET"])
def api_dsag_list_cached():
    """Return a list of all file-cached graph configurations."""
    try:
        cached = list_cached_graphs()
        return jsonify({"success": True, "cached_configs": cached})
    except Exception as exc:
        return jsonify({"success": False, "error": str(exc)})


@app.route("/api/dsag/load_cached", methods=["POST"])
def api_dsag_load_cached():
    """
    Load a previously cached graph by cache_key.

    Request body: { "cache_key": "..." }
    """
    try:
        data = request.get_json()
        cache_key = (data.get("cache_key") or "").strip()
        if not cache_key:
            return jsonify({"success": False, "error": "cache_key is required"})

        state = load_graph_from_file(cache_key)
        if not state or not state.is_ready():
            return jsonify({"success": False, "error": "Cache not found or invalid"})

        set_dsag_state(state, cache_key)

        # Restore session context from graph metadata
        topic = ""
        if state.graph:
            topic = state.graph.topic or ""
            session["dsag_topic"] = topic
            session["dsag_researcher_bg"] = state.graph.researcher_bg or ""
            session["dsag_expert_bg"] = state.graph.expert_bg or ""
        elif state.question_graphs:
            g = state.question_graphs[0].graph
            topic = g.topic or ""
            session["dsag_topic"] = topic
            session["dsag_researcher_bg"] = g.researcher_bg or ""
            session["dsag_expert_bg"] = g.expert_bg or ""

        # Export artifacts
        if state.graph:
            export_graph_artifacts(state.graph)
        export_multi_graph_artifacts_from_state(state)

        # Transcript summary
        if state.transcript_summary:
            set_transcript_summary(state.transcript_summary)

        return jsonify({
            "success": True,
            "cached": True,
            "cache_key": cache_key,
            "cache_source": "file",
            "metadata": (
                state.graph.metadata
                if state.graph
                else {"question_graph_count": len(state.question_graphs)}
            ),
        })
    except Exception as exc:
        return jsonify({"success": False, "error": str(exc)})


@app.route("/api/dsag/status", methods=["GET"])
def api_dsag_status():
    """
    Get the status of the current DSAG session.
    
    Response:
    {
        "ready": true/false,
        "status": "building|ready|error",
        "error": "...",
        "metadata": {...}
    }
    """
    try:
        state = get_dsag_state()
        if state is None:
            return jsonify({
                "ready": False,
                "status": "not_initialized",
                "error": "DSAG not initialized. Call /api/dsag/init first.",
            })
        
        return jsonify({
            "ready": state.is_ready(),
            "status": state.status,
            "error": state.error,
            "metadata": (
                state.graph.metadata
                if state.graph
                else {"question_graph_count": len(state.question_graphs)}
            ),
        })
    
    except Exception as exc:
        return jsonify({"ready": False, "status": "error", "error": str(exc)})


@app.route("/api/dsag/analyze_turn", methods=["POST"])
def api_dsag_analyze_turn():
    """
    Analyze a conversation turn using the DSAG graph.

    Request body:
    {
        "researcher_question": "The researcher's question",
        "expert_answer": "The expert's response"
    }

    Response:
    {
        "success": true,
        "analysis": {
            "located": {...},
            "divergence": {...},
            "assistance": {
                "relation_type": "LexicalGap|ConceptualGap|TacitGap|ScopeGap|ProcessGap",
                "payload": {...}
            },
            ...
        }
    }
    """
    try:
        state = get_dsag_state()
        if state is None or not state.is_ready():
            return jsonify({
                "success": False,
                "error": "DSAG not ready. Call /api/dsag/init first.",
            })

        data = request.get_json()
        researcher_question = data.get("researcher_question", "").strip()
        expert_answer = data.get("expert_answer", "").strip()

        if not expert_answer:
            return jsonify({"success": False, "error": "Expert answer is required"})

        analysis = analyze_turn_with_dsag(
            state,
            researcher_question,
            expert_answer,
            get_messages(),
        )

        return jsonify({
            "success": True,
            "analysis": analysis.to_dict(),
        })

    except Exception as exc:
        return jsonify({"success": False, "error": str(exc)})


@app.route("/api/dsag/graph", methods=["GET"])
def api_dsag_graph():
    """
    Get the full DSAG graph structure (for debugging/visualization).
    
    Response:
    {
        "success": true,
        "graph": {...}
    }
    """
    try:
        state = get_dsag_state()
        if state is None or not state.is_ready():
            return jsonify({
                "success": False,
                "error": "DSAG not ready. Call /api/dsag/init first.",
            })
        
        return jsonify({
            "success": True,
            "graph": state.graph.to_dict() if state.graph else None,
            "question_graphs": [
                {
                    "question_id": qg.question_id,
                    "question_text": qg.question_text,
                    "depends_on": qg.depends_on,
                    "graph": qg.graph.to_dict(),
                }
                for qg in state.question_graphs
            ],
        })
    
    except Exception as exc:
        return jsonify({"success": False, "error": str(exc)})


# ============== Main Routes ==============

@app.route("/", methods=["GET", "POST"])
def index():
    # Clear session if requested via query param
    if request.args.get("reset") == "1":
        clear_dsag_state()
        session.clear()
        return redirect(url_for("index"))

    messages = get_messages()
    if request.method == "POST":
        action = request.form.get("action", "chat")

        if action == "save_guide":
            guide_text = request.form.get("guide_text", "")
            set_guide_text(guide_text)
            set_guide_error(None)
            return redirect(url_for("index"))

        if action == "upload_guide":
            uploaded = request.files.get("guide_file")
            if not uploaded or not uploaded.filename:
                set_guide_error("Please choose a file to upload.")
            else:
                try:
                    extracted = extract_text_from_upload(uploaded)
                    set_guide_text(extracted)
                    set_guide_error(None)
                except Exception as exc:
                    set_guide_error(str(exc))
            return redirect(url_for("index"))

        if action == "upload_questionnaire":
            uploaded = request.files.get("questionnaire_file")
            if not uploaded or not uploaded.filename:
                set_questionnaire_error("Please choose an interview script file to upload.")
            else:
                try:
                    extracted = extract_text_from_upload(uploaded).strip()
                    if not extracted:
                        raise RuntimeError("Uploaded interview script is empty after extraction.")
                    set_uploaded_questionnaire_text(extracted)
                    set_questionnaire_error(None)
                except Exception as exc:
                    set_questionnaire_error(str(exc))
            return redirect(url_for("index"))

        if action == "toggle_mismap":
            try:
                msg_index = int(request.form.get("msg_index", "-1"))
            except ValueError:
                msg_index = -1
            if 0 <= msg_index < len(messages):
                target = messages[msg_index]
                if target.get("role") == "expert":
                    if target.get("manual_mismap"):
                        clear_manual_mismap(target)
                    else:
                        ensure_manual_mismap(target)
                session["messages"] = messages
            return redirect(url_for("index"))

        # Handle chat messages
        researcher_text = request.form.get("researcher_input", "").strip()
        expert_text = request.form.get("expert_input", "").strip()
        researcher_source = request.form.get("researcher_source", "text")
        expert_source = request.form.get("expert_source", "text")

        ts_for_turns = get_transcript_summary()
        pending_turn_index = (
            ts_for_turns.last_updated_turn + 1
            if ts_for_turns and ts_for_turns.main_bullets
            else None
        )

        if researcher_text:
            researcher_msg = {
                "role": "researcher",
                "content": researcher_text,
                "source": researcher_source,
            }
            if pending_turn_index is not None:
                researcher_msg["turn_index"] = pending_turn_index
            messages.append(researcher_msg)

        if researcher_text and not expert_text:
            topic = str(session.get("dsag_topic", "")).strip()
            demographics = str(session.get("interviewee_demographics", "")).strip()
            interviewer_demographics = str(session.get("dsag_researcher_bg", "")).strip()
            history = build_context_summary(messages[:-1], max_turns=3)
            try:
                expert_text = generate_ai_interviewee_reply(
                    question=researcher_text,
                    topic=topic,
                    demographics=demographics,
                    interviewer_demographics=interviewer_demographics,
                    history=history,
                )
                expert_source = "ai_interviewee"
            except Exception as interviewee_exc:
                print(f"[AI interviewee] Error: {interviewee_exc}")

        if expert_text:
            msg = {"role": "expert", "content": expert_text, "source": expert_source}
            turn_index = pending_turn_index
            last_question = ""
            for m in reversed(messages):
                if m["role"] == "researcher":
                    last_question = m["content"]
                    if turn_index is None and m.get("turn_index") is not None:
                        turn_index = m.get("turn_index")
                    break
            if turn_index is not None:
                msg["turn_index"] = turn_index

            # ---- DSAG auto-analysis ----
            dsag_state = get_dsag_state()
            if dsag_state and dsag_state.is_ready() and last_question:
                try:
                    dsag_analysis = analyze_turn_with_dsag(
                        dsag_state,
                        last_question,
                        expert_text,
                        messages,
                    )

                    # Store analysis result in message for template rendering
                    msg["dsag_analysis"] = dsag_analysis.to_dict()
                except Exception as dsag_exc:
                    print(f"[DSAG auto-analysis] Error: {dsag_exc}")

            # Lightweight per-message term annotation for inline highlighting.
            # This runs independently of the selected DSAG gap type.
            try:
                annotation_context = build_term_annotation_context(last_question)
                analysis_result = analyze_exchange(last_question, expert_text, context=annotation_context)
                jargon_terms = analysis_result.get("jargon", []) if isinstance(analysis_result, dict) else []
                if isinstance(jargon_terms, list):
                    msg["jargon_terms"] = jargon_terms
            except Exception as annotation_exc:
                print(f"[Term annotation] Error: {annotation_exc}")

            messages.append(msg)

        session["messages"] = messages
        return redirect(url_for("index"))

    # Check if DSAG is ready for template rendering
    dsag_state = get_dsag_state()
    dsag_ready = dsag_state is not None and dsag_state.is_ready()

    # Build transcript summary for template
    ts = get_transcript_summary()
    ts_data = ts.to_dict() if ts else None

    return render_template(
        "index.html",
        messages=messages,
        guide_text=get_guide_text(),
        guide_error=session.get("guide_error"),
        questionnaire_uploaded=bool(get_uploaded_questionnaire_text().strip()),
        questionnaire_text=get_uploaded_questionnaire_text(),
        questionnaire_error=session.get("questionnaire_error"),
        dsag_ready=dsag_ready,
        transcript_summary=ts_data,
    )


if __name__ == "__main__":
    # Disable debug reloader to keep interview sessions stable during runtime.
    app.run(debug=False, use_reloader=False)
